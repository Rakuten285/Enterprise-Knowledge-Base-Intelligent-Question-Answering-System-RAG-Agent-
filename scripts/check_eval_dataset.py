"""
评测集自检脚本

对每条 Q&A 做两项机械检查：
  1. reference_contexts 里的每段文字能否在对应原文中精确找到
  2. ground_truth 里出现的所有数字能否在对应原文中找到

输出：
  - PASS：两项都通过
  - WARN：数字在原文里找不到（可能是推断或错误）
  - FAIL：reference_contexts 原文里找不到（直接引用错误）

用法：
    python scripts/check_eval_dataset.py              # 检查全部
    python scripts/check_eval_dataset.py --id 7 8 9  # 只检查指定id
    python scripts/check_eval_dataset.py --fail-only  # 只显示FAIL和WARN
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DATASET_PATH = ROOT / "data" / "eval_dataset.json"
RAW_DIR = ROOT / "data" / "raw"


def load_source(source_doc: str) -> str:
    """读取原始文档内容（纯文本）"""
    path = RAW_DIR / source_doc
    if not path.exists():
        return ""
    if path.suffix == ".docx":
        try:
            from docx import Document
            doc = Document(path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
            return "\n".join(parts)
        except Exception:
            return ""
    return path.read_text(encoding="utf-8")


def extract_numbers(text: str) -> list[str]:
    """提取文本中的数字（含小数、百分比、带单位的数）"""
    return re.findall(r'\d+(?:\.\d+)?(?:\s*[%元天年月日小时万])?', text)


def normalize(text: str) -> str:
    """去掉空白和 Markdown 标记后用于子串比较"""
    text = re.sub(r'[*_#`>|]', '', text)   # 去掉 Markdown 符号
    text = re.sub(r'\s+', '', text)          # 去掉所有空白
    return text

def check_context_in_source(context: str, source_text: str) -> bool:
    """检查 context 是否能在原文中找到（忽略空白和 Markdown 标记）"""
    return normalize(context) in normalize(source_text)


def check_numbers_in_source(ground_truth: str, source_text: str) -> list[str]:
    """返回 ground_truth 中在原文里找不到的数字"""
    src_norm = normalize(source_text)
    missing = []
    for num in extract_numbers(ground_truth):
        if normalize(num) not in src_norm:
            missing.append(num)
    return missing


def check_sample(sample: dict, source_cache: dict) -> dict:
    sid = sample["id"]
    source_doc = sample["source_doc"]

    if source_doc not in source_cache:
        source_cache[source_doc] = load_source(source_doc)
    source_text = source_cache[source_doc]

    if not source_text:
        return {"id": sid, "status": "ERROR", "reason": f"找不到原文文件: {source_doc}"}

    # 检查1：reference_contexts 原文匹配
    ctx_failures = []
    for ctx in sample.get("reference_contexts", []):
        if not check_context_in_source(ctx, source_text):
            ctx_failures.append(ctx[:50] + "...")

    # 检查2：ground_truth 数字验证
    missing_nums = check_numbers_in_source(sample["ground_truth"], source_text)

    if ctx_failures:
        status = "FAIL"
        reason = f"reference_contexts 在原文中找不到: {ctx_failures}"
    elif missing_nums:
        status = "WARN"
        reason = f"ground_truth 中的数字在原文找不到: {missing_nums}"
    else:
        status = "PASS"
        reason = ""

    return {"id": sid, "status": status, "reason": reason,
            "question": sample["question"], "source_doc": source_doc}


def main():
    parser = argparse.ArgumentParser(description="评测集自检")
    parser.add_argument("--id", nargs="+", type=int, help="只检查指定id")
    parser.add_argument("--fail-only", action="store_true", help="只显示FAIL和WARN")
    args = parser.parse_args()

    dataset = json.loads(DATASET_PATH.read_text(encoding="utf-8"))

    if args.id:
        dataset = [s for s in dataset if s["id"] in args.id]

    source_cache = {}
    results = [check_sample(s, source_cache) for s in dataset]

    pass_count = sum(1 for r in results if r["status"] == "PASS")
    warn_count = sum(1 for r in results if r["status"] == "WARN")
    fail_count = sum(1 for r in results if r["status"] in ("FAIL", "ERROR"))

    print(f"\n{'='*60}")
    print(f"评测集自检结果：共 {len(results)} 条")
    print(f"  PASS: {pass_count}  WARN: {warn_count}  FAIL: {fail_count}")
    print(f"{'='*60}\n")

    for r in results:
        if args.fail_only and r["status"] == "PASS":
            continue
        tag = {"PASS": "OK", "WARN": "!!", "FAIL": "XX", "ERROR": "EE"}.get(r["status"], "?")
        print(f"[{tag}] id={r['id']:3d}  {r['status']}  {r['source_doc']}")
        print(f"      问题: {r['question']}")
        if r["reason"]:
            print(f"      原因: {r['reason']}")
        print()

    if fail_count > 0 or warn_count > 0:
        sys.exit(1)


if __name__ == "__main__":
    main()
