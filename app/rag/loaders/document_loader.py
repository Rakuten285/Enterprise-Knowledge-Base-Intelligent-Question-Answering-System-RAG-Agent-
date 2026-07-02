"""
文档加载模块
支持 PDF / Word(.docx) / TXT / Markdown，统一返回 LangChain Document 列表
每个 Document.metadata 至少包含 source / file_type，PDF 额外包含 page
"""
from __future__ import annotations

from pathlib import Path
from typing import Callable

from langchain_core.documents import Document

from app.core.logging import logger

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}


def load_pdf(file_path: Path) -> list[Document]:
    """逐页解析 PDF，保留页码信息，便于后续追溯引用来源"""
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    docs: list[Document] = []
    for page_idx, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if not text:
            continue
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": file_path.name,
                    "file_type": "pdf",
                    "page": page_idx + 1,
                },
            )
        )
    if not docs:
        logger.warning(f"PDF 未提取到任何文本，可能是扫描件未做 OCR: {file_path}")
    return docs


def load_docx(file_path: Path) -> list[Document]:
    """解析 Word 文档，按段落聚合，过滤空段落；表格内容单独提取拼接"""
    import docx

    document = docx.Document(str(file_path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    table_texts: list[str] = []
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_texts.append(row_text)

    full_text = "\n".join(paragraphs)
    if table_texts:
        full_text += "\n\n[表格内容]\n" + "\n".join(table_texts)

    if not full_text.strip():
        logger.warning(f"Word 文档未提取到任何文本: {file_path}")
        return []

    return [
        Document(
            page_content=full_text,
            metadata={"source": file_path.name, "file_type": "docx"},
        )
    ]


def load_txt(file_path: Path) -> list[Document]:
    text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        return []
    file_type = "md" if file_path.suffix.lower() == ".md" else "txt"
    return [
        Document(
            page_content=text,
            metadata={"source": file_path.name, "file_type": file_type},
        )
    ]


_LOADER_REGISTRY: dict[str, Callable[[Path], list[Document]]] = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".doc": load_docx,
    ".txt": load_txt,
    ".md": load_txt,
}


def load_document(file_path: str | Path) -> list[Document]:
    """根据文件后缀分发到对应解析器"""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    suffix = path.suffix.lower()
    loader = _LOADER_REGISTRY.get(suffix)
    if loader is None:
        raise ValueError(
            f"不支持的文件类型: {suffix}，目前支持 {sorted(SUPPORTED_EXTENSIONS)}"
        )

    logger.info(f"加载文档: {path.name} (类型={suffix})")
    docs = loader(path)
    logger.info(f"解析完成: {path.name} -> {len(docs)} 个原始片段（页/段）")
    return docs


def load_documents_from_dir(dir_path: str | Path) -> list[Document]:
    """批量加载目录下所有支持的文档"""
    dir_path = Path(dir_path)
    all_docs: list[Document] = []
    files = [
        f
        for f in sorted(dir_path.rglob("*"))
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    logger.info(f"在 {dir_path} 中发现 {len(files)} 个可解析文件")
    for f in files:
        try:
            all_docs.extend(load_document(f))
        except Exception as e:
            logger.error(f"解析失败，跳过文件 {f.name}: {e}")
    return all_docs
